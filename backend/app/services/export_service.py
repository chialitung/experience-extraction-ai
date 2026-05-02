import io
from typing import Dict, Any
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import markdown as md


class ExportService:
    """成果导出服务

    支持将访谈成果导出为多种格式：Markdown、Word(docx)、PDF。
    """

    @classmethod
    def export_markdown(cls, content: str, filename_base: str) -> tuple:
        """导出为Markdown格式
        
        Returns: (bytes, filename)
        """
        data = content.encode("utf-8")
        filename = f"{filename_base}.md"
        return data, filename

    @classmethod
    def export_docx(cls, content: str, title: str, filename_base: str) -> tuple:
        """导出为Word文档格式

        Markdown → HTML（markdown库） → DOCX（BeautifulSoup遍历HTML树）
        Returns: (bytes, filename)
        """
        from bs4 import BeautifulSoup

        def _process_table(doc, table_element):
            rows = table_element.find_all('tr')
            if not rows:
                return
            max_cols = max(len(row.find_all(['td', 'th'])) for row in rows)
            if max_cols == 0:
                return
            table = doc.add_table(rows=len(rows), cols=max_cols)
            table.style = 'Table Grid'
            for i, row in enumerate(rows):
                cells = row.find_all(['td', 'th'])
                for j, cell in enumerate(cells):
                    if j < max_cols:
                        table.rows[i].cells[j].text = cell.get_text()
                        if cell.name == 'th':
                            for paragraph in table.rows[i].cells[j].paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True

        def _process_element(doc, element, paragraph=None, list_level=0):
            for child in element.children:
                if isinstance(child, str):
                    text = str(child)
                    if text and paragraph:
                        paragraph.add_run(text)
                    continue

                if child.name in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
                    level = int(child.name[1])
                    doc.add_heading(child.get_text(), level=min(level, 4))
                    paragraph = None

                elif child.name == 'p':
                    if paragraph:
                        _process_element(doc, child, paragraph)
                    else:
                        p = doc.add_paragraph()
                        _process_element(doc, child, p)
                    paragraph = None

                elif child.name == 'blockquote':
                    p = doc.add_paragraph()
                    p.style = 'Quote'
                    _process_element(doc, child, p)
                    paragraph = None

                elif child.name == 'ul':
                    for li in child.find_all('li', recursive=False):
                        bullet_styles = ['List Bullet', 'List Bullet 2', 'List Bullet 3']
                        style = bullet_styles[min(list_level, 2)]
                        p = doc.add_paragraph(style=style)
                        _process_element(doc, li, p, list_level + 1)
                    paragraph = None

                elif child.name == 'ol':
                    for li in child.find_all('li', recursive=False):
                        number_styles = ['List Number', 'List Number 2', 'List Number 3']
                        style = number_styles[min(list_level, 2)]
                        p = doc.add_paragraph(style=style)
                        _process_element(doc, li, p, list_level + 1)
                    paragraph = None

                elif child.name == 'table':
                    _process_table(doc, child)
                    paragraph = None

                elif child.name == 'pre':
                    p = doc.add_paragraph()
                    p.style = 'Intense Quote'
                    code = child.find('code')
                    if code:
                        run = p.add_run(code.get_text())
                    else:
                        run = p.add_run(child.get_text())
                    run.font.name = 'Courier New'
                    paragraph = None

                elif child.name == 'br':
                    if paragraph:
                        paragraph.add_run('\n')

                elif child.name in ('strong', 'b'):
                    if paragraph:
                        run = paragraph.add_run(child.get_text())
                        run.bold = True

                elif child.name in ('em', 'i'):
                    if paragraph:
                        run = paragraph.add_run(child.get_text())
                        run.italic = True

                elif child.name == 'code':
                    if paragraph:
                        run = paragraph.add_run(child.get_text())
                        run.font.name = 'Courier New'
                        run.font.size = Pt(10)

                elif child.name == 'a':
                    if paragraph:
                        run = paragraph.add_run(child.get_text())
                        run.font.color.rgb = RGBColor(0, 102, 204)
                        run.underline = True

                elif child.name in ('span', 'div', 'li'):
                    _process_element(doc, child, paragraph, list_level)

                else:
                    if paragraph:
                        paragraph.add_run(child.get_text())
                    else:
                        p = doc.add_paragraph()
                        p.add_run(child.get_text())

        doc = Document()

        # 设置所有基础样式字体为微软雅黑
        for style_name in ['Normal'] + [f'Heading {i}' for i in range(1, 10)]:
            if style_name in doc.styles:
                style = doc.styles[style_name]
                style.font.name = 'Microsoft YaHei'
                style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

        style = doc.styles['Normal']
        style.font.size = Pt(11)

        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        time_para = doc.add_paragraph()
        time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = time_para.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph()

        # Markdown → HTML
        html_content = md.markdown(
            content,
            extensions=['tables', 'fenced_code', 'toc'],
        )
        soup = BeautifulSoup(f'<html><body>{html_content}</body></html>', 'html.parser')

        # HTML → DOCX
        body = soup.find('body')
        if body:
            _process_element(doc, body)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        filename = f"{filename_base}.docx"
        return buffer.getvalue(), filename

    @classmethod
    def export_pdf(cls, content: str, title: str, filename_base: str) -> tuple:
        """导出为PDF格式

        使用 WeasyPrint 将 Markdown 渲染的 HTML 转为 PDF 文件。

        Returns: (bytes, filename)
        """
        from weasyprint import HTML

        html_content = md.markdown(
            content,
            extensions=['tables', 'fenced_code', 'toc'],
        )

        full_html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<title>{title}</title>
<style>
@page {{ size: A4; margin: 2cm; }}
body {{ font-family: "Microsoft YaHei", "PingFang SC", sans-serif; line-height: 1.8; color: #333; }}
h1 {{ color: #1a1a1a; border-bottom: 2px solid #4f46e5; padding-bottom: 10px; }}
h2 {{ color: #333; margin-top: 30px; border-left: 4px solid #4f46e5; padding-left: 12px; }}
h3 {{ color: #444; margin-top: 20px; }}
blockquote {{ border-left: 4px solid #ddd; padding-left: 16px; color: #666; margin: 0; }}
code {{ background: #f4f4f4; padding: 2px 6px; border-radius: 3px; font-family: monospace; }}
pre {{ background: #f4f4f4; padding: 16px; border-radius: 6px; overflow-x: auto; }}
table {{ border-collapse: collapse; width: 100%; margin: 16px 0; }}
th, td {{ border: 1px solid #ddd; padding: 8px 12px; text-align: left; }}
th {{ background: #f8f8f8; }}
ul, ol {{ padding-left: 20px; }}
</style>
</head>
<body>
{html_content}
</body>
</html>"""

        buffer = io.BytesIO()
        HTML(string=full_html).write_pdf(buffer)
        buffer.seek(0)
        filename = f"{filename_base}.pdf"
        return buffer.getvalue(), filename

    @classmethod
    def export_json(cls, data: Dict[str, Any], filename_base: str) -> tuple:
        """导出为JSON格式
        
        Returns: (bytes, filename)
        """
        import json
        content = json.dumps(data, ensure_ascii=False, indent=2)
        data_bytes = content.encode("utf-8")
        filename = f"{filename_base}.json"
        return data_bytes, filename
